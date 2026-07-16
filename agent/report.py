"""Export an analysis as a regulated-style Statistical Analysis Report (.docx).

Renders an AgentResult into an industry-format Word document with the sections a reviewer of a
clinical Statistical Analysis Report (SAR) expects: a title/approval page with an explicit DRAFT
status and signature block, a synopsis, data sources and analysis populations, the statistical
methods with the exact software environment (versions pulled at runtime), numbered and captioned
results tables/figures with methodological footnotes, assumption diagnostics, conclusions, and a
limitations/validation statement — with a confidential, page-numbered footer throughout.

Honest by construction: it documents that this is an automated EXPLORATORY analysis on synthetic
data, not a validated, pre-specified, double-programmed regulatory submission.
"""
from __future__ import annotations

import datetime as _dt
import io

# One-line, reviewer-facing description of each method (drives the Methods section).
_METHOD_BLURB = {
    "logistic": "Multivariable logistic regression estimating adjusted odds ratios with 95% "
                "confidence intervals (Wald), fit by maximum likelihood.",
    "ols": "Multivariable ordinary-least-squares linear regression estimating adjusted "
           "coefficients with 95% confidence intervals.",
    "cox": "Cox proportional-hazards regression estimating adjusted hazard ratios with 95% "
           "confidence intervals; the proportional-hazards assumption was assessed.",
    "survival": "Kaplan–Meier survival estimation with Cox proportional-hazards modelling for "
                "adjusted hazard ratios.",
    "forest": "Random-forest model with 5-fold cross-validated discrimination and permutation "
              "importance computed on a held-out split (imputation fit inside the pipeline).",
    "timeseries": "Holt–Winters exponential smoothing with an approximate, horizon-widening "
                  "prediction band.",
    "causal": "Cross-fitted AIPW doubly-robust estimator (potential-outcome + propensity forests) for the "
              "average treatment effect, with an influence-function 95% CI; observational, not randomized.",
    "experiment": "Two-group experiment analysis: per-arm rates/means with confidence intervals, "
                  "Newcombe/Welch interval on the difference, and Benjamini–Hochberg FDR control "
                  "across variants.",
    "noninferiority": "Non-inferiority assessment: Farrington–Manning score test (binary) or the "
                      "confidence-interval–versus–margin rule (continuous) at one-sided α = 0.025.",
    "sample_size": "Design-stage sample-size / power calculation (normal approximation for "
                   "proportions, t-test power for means).",
    "association": "Descriptive association summary with confidence intervals.",
    "assurance": "Bayesian design-stage go/no-go. Assurance (the probability of a GO, averaged over "
                 "the prior uncertainty about the true effect) with a dual-criterion decision rule "
                 "(Target Value / Lower Reference Value), a prior-sensitivity panel, and simulated "
                 "operating characteristics. Conjugate Beta-Binomial; computed in closed form.",
    "interim": "Bayesian interim go/no-go. Single-arm: posterior response rate with a 95% credible "
               "interval and the exact predictive probability the trial ends in a GO at full enrolment. "
               "Two-arm: the same decision on the risk difference (treatment - control) via an exact "
               "joint beta-binomial predictive. Conjugate Beta-Binomial; closed form (no simulation).",
}


def _kv(doc, key, val):
    p = doc.add_paragraph()
    p.add_run(f"{key}: ").bold = True
    p.add_run(str(val))
    return p


def _split_notes(issues):
    """Separate data-preparation steps from model/assumption diagnostics for the report."""
    prep, diag = [], []
    for i in issues or []:
        low = i.lower()
        if low.startswith(("dropped", "imputed", "removed")):
            prep.append(i)
        else:
            diag.append(i)
    return prep, diag


def _mono(p):
    from docx.shared import Pt
    for r in p.runs:
        r.font.name = "Consolas"
        r.font.size = Pt(9)


def _footnote(doc, text):
    """A small italic methodological footnote under a table or figure."""
    from docx.shared import Pt
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.size = Pt(8)
    return p


def _versions() -> dict:
    """The exact runtime software environment, pulled at generation time (not hard-coded)."""
    import platform
    from importlib.metadata import PackageNotFoundError, version
    vers = {"Python": platform.python_version()}
    for dist in ("statsmodels", "scikit-learn", "scipy", "numpy", "pandas", "duckdb"):
        try:
            vers[dist] = version(dist)
        except PackageNotFoundError:
            vers[dist] = "not installed"
        except Exception:  # noqa: BLE001
            vers[dist] = "—"
    return vers


def _field(paragraph, instr: str):
    """Append a Word field (e.g. PAGE, NUMPAGES) to a paragraph — computed by Word on open."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar"); begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText"); instr_el.set(qn("xml:space"), "preserve"); instr_el.text = instr
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for el in (begin, instr_el, sep, end):
        run._r.append(el)


def _confidential_footer(doc):
    """A centered, page-numbered confidentiality footer on every page."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt
    footer = doc.sections[0].footer
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Confidential — automated draft, not for regulatory submission     Page ")
    r.font.size = Pt(8)
    _field(p, "PAGE")
    r2 = p.add_run(" of "); r2.font.size = Pt(8)
    _field(p, "NUMPAGES")
    for run in p.runs:
        run.font.size = Pt(8)


def _sig_line(doc, role):
    from docx.shared import Pt
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)          # tight, so all three lines fit on the title page
    p.add_run(f"{role}:").bold = True
    p.add_run("  ____________________  Name/Title: ______________  Date: __________")
    for r in p.runs:
        r.font.size = Pt(9.5)


def _chart_png(chart):
    """Render an Altair chart to PNG bytes (vl-convert); None if it can't be rendered.
    Rendered at 200 ppi for a crisp figure on the printed page."""
    if chart is None:
        return None
    try:
        buf = io.BytesIO()
        chart.save(buf, format="png", ppi=200)
        return buf.getvalue()
    except Exception:
        return None


def _cell(v) -> str:
    """Human-readable table cell: NaN / None / inf render as an em dash, not literal 'nan'/'None'/'inf'."""
    import math
    if v is None:
        return "—"
    if isinstance(v, float) and (math.isnan(v) or math.isinf(v)):
        return "—"
    return str(v)


def _figure_titles(result, m):
    """(figure, caption) pairs matching the on-screen result, so the report shows the same charts."""
    from agent import charts as ch
    mt = m.get("model_type")
    pairs = []
    try:
        if mt == "timeseries":
            pairs = [(ch.forecast_chart(m.get("series")), "Observed history and forecast with 95% band.")]
        elif mt == "forest":
            pairs = [(ch.importance_chart(m), "Permutation feature importance (held-out split).")]
        elif mt == "experiment":
            pairs = [(ch.experiment_chart(m), "Per-arm estimates with 95% confidence intervals.")]
        elif mt == "noninferiority":
            pairs = [(ch.ni_plot(m), "Treatment−control difference versus the non-inferiority margin.")]
        elif mt == "sample_size":
            pairs = [(ch.power_curve_chart(m), "Required sample size across target power.")]
        elif mt == "survival":
            pairs = [(ch.survival_plot(m.get("km")), "Kaplan–Meier survival curves."),
                     (ch.forest_plot(m), "Adjusted hazard ratios with 95% confidence intervals.")]
        elif m.get("terms"):                         # logistic / ols / cox / causal / association
            pairs = [(ch.forest_plot(m), f"Adjusted {m.get('effect_label', 'estimates')} with 95% CIs.")]
        elif not m and getattr(result, "dataframe", None) is not None:
            pairs = [(ch.build_chart(result.dataframe), "Summary of the query result.")]
    except Exception:
        pairs = []
    return [(f, c) for f, c in pairs if f is not None]


def build_docx(result, *, when: _dt.datetime | None = None) -> bytes:
    """Build the industry-format .docx SAR for an AgentResult and return its bytes."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    now = when or _dt.datetime.now()
    ts = now.strftime("%d %b %Y %H:%M")
    doc_id = f"CIA-SAR-{now.strftime('%Y%m%d-%H%M')}"
    m = result.model or {}
    mt = m.get("model_type")

    doc = Document()
    _confidential_footer(doc)

    # counters for numbered captions
    counters = {"table": 0, "figure": 0}

    def table_caption(title):
        counters["table"] += 1
        p = doc.add_paragraph()
        p.add_run(f"Table {counters['table']}. ").bold = True
        p.add_run(title)
        return p

    def figure_caption(title):
        counters["figure"] += 1
        p = doc.add_paragraph()
        r = p.add_run(f"Figure {counters['figure']}. ")
        r.bold = True; r.italic = True
        p.add_run(title).italic = True
        return p

    # ───────────────────────── title / approval page ─────────────────────────
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = t.add_run("STATISTICAL ANALYSIS REPORT"); tr.bold = True; tr.font.size = Pt(22)
    st = doc.add_paragraph(); st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = st.add_run("Automated exploratory analysis · Clinical Insight Agent")
    sr.italic = True; sr.font.size = Pt(12)

    status = doc.add_paragraph(); status.alignment = WD_ALIGN_PARAGRAPH.CENTER
    stat_r = status.add_run("STATUS: DRAFT — FOR INTERNAL REVIEW ONLY")
    stat_r.bold = True; stat_r.font.size = Pt(13); stat_r.font.color.rgb = RGBColor(0xB0, 0x00, 0x00)

    doc.add_paragraph()
    meta = doc.add_table(rows=0, cols=2); meta.style = "Table Grid"
    for k, v in [
        ("Document identifier", doc_id),
        ("Version", "0.1 (draft)"),
        ("Report date", ts),
        ("Analysis engine", "Clinical Insight Agent (LLM: OpenAI)"),
        ("Sponsor / Study", "________________________  (to be completed)"),
        ("Analysis type", "Exploratory — hypothesis-generating"),
        ("Data", "Synthetic / illustrative — not clinical fact"),
    ]:
        c = meta.add_row().cells
        c[0].text = k; c[0].paragraphs[0].runs[0].bold = True
        c[1].text = v

    ce = doc.add_paragraph()
    ce.add_run("Confidentiality. ").bold = True
    ce.add_run("This document contains automated analytical output and is intended for internal "
               "review only. It is a machine-generated draft and must not be used for a regulatory "
               "submission or clinical decision without qualified biostatistical review.")

    _ps = m.get("prespec") or {}
    if _ps.get("status"):
        _kv(doc, "Pre-specification", _ps["status"])
        if _ps.get("drift"):
            doc.add_paragraph("Departures from the locked design: "
                              + "; ".join(f"{d['field']}: locked {d['locked']} -> used {d['actual']}"
                                          for d in _ps["drift"]))
        if (_ps.get("lock") or {}).get("lock_id"):     # interim without a lock stores lock=None
            _kv(doc, "Design lock", _ps["lock"]["lock_id"][:16])

    doc.add_heading("Approval / review", 2)
    doc.add_paragraph("This draft requires review and sign-off before use:")
    _sig_line(doc, "Prepared by")
    _sig_line(doc, "Reviewed by (Biostatistics)")
    _sig_line(doc, "Approved by")
    doc.add_page_break()

    # ───────────────────────── 1. synopsis ─────────────────────────
    doc.add_heading("1. Synopsis", 1)
    verdict = m.get("verdict") or {}
    syn = doc.add_paragraph()
    syn.add_run("Objective. ").bold = True
    syn.add_run((result.question or "—") + "  ")
    if mt:
        syn.add_run("Method. ").bold = True
        syn.add_run(_METHOD_BLURB.get(mt, mt.upper()) + "  ")
    if verdict.get("call"):
        syn.add_run("Primary conclusion. ").bold = True
        syn.add_run(f"{verdict['call']}. {verdict.get('reason', '')}")

    # ───────────────────────── 2. objective ─────────────────────────
    doc.add_heading("2. Objective", 1)
    doc.add_paragraph(result.question or "")
    if getattr(result, "hypothesis", None):
        doc.add_heading("2.1 Hypothesis", 2)
        doc.add_paragraph(result.hypothesis)

    # ───────────────────────── 3. data sources & populations ─────────────────────────
    doc.add_heading("3. Data sources and analysis population", 1)
    if getattr(result, "citations", None):
        _kv(doc, "Source tables", ", ".join(result.citations))
    if m.get("n") is not None:
        _kv(doc, "Analysis set (n analyzed)", f"{m.get('n', 0):,}")
    doc.add_paragraph(
        "Analysis population: complete-case after the data-engineering steps in §3.2. Variable "
        "selection was data-driven (removal of quasi-constant, high-missingness, high-cardinality, and "
        "collinear predictors) — appropriate for exploration but NOT for a confirmatory analysis, where "
        "the analysis set and covariates must be pre-specified in a Statistical Analysis Plan.")
    if getattr(result, "sql", None):
        doc.add_heading("3.1 Analytic query", 2)
        _mono(doc.add_paragraph(result.sql))

    prep, diag = _split_notes(m.get("issues"))
    if prep:
        doc.add_heading("3.2 Data preparation (audit trail)", 2)
        for s in prep:
            doc.add_paragraph(s, style="List Bullet")

    # ───────────────────────── 4. statistical methods ─────────────────────────
    doc.add_heading("4. Statistical methods", 1)
    if mt:
        doc.add_paragraph(_METHOD_BLURB.get(mt, mt.upper()))
    if m.get("fit_stat"):
        _kv(doc, "Fit / design", m["fit_stat"])
    if m.get("note"):
        doc.add_paragraph(m["note"])
    doc.add_heading("4.1 Software environment", 2)
    table_caption("Software versions used to generate this analysis.")
    vt = doc.add_table(rows=1, cols=2); vt.style = "Table Grid"
    vt.rows[0].cells[0].text = "Component"; vt.rows[0].cells[1].text = "Version"
    for k, v in _versions().items():
        c = vt.add_row().cells; c[0].text = k; c[1].text = v
    _footnote(doc, "Versions captured at report-generation time from the running environment.")

    # ───────────────────────── 5. results ─────────────────────────
    doc.add_heading("5. Results", 1)
    if verdict:
        _kv(doc, "Conclusion", verdict.get("call", ""))
        if verdict.get("reason"):
            doc.add_paragraph(verdict["reason"])

    if mt == "sample_size":
        table_caption("Required sample size by arm.")
        st = doc.add_table(rows=1, cols=2); st.style = "Table Grid"
        st.rows[0].cells[0].text = "Arm"; st.rows[0].cells[1].text = "Subjects required"
        for a in m.get("arms", []):
            c = st.add_row().cells; c[0].text = str(a["arm"]); c[1].text = f"{a['n']:,}"
        _footnote(doc, "Analytic power; inflate for anticipated dropout / non-evaluable subjects.")
    elif mt == "timeseries":                          # forecast output, not the raw input series
        series = m.get("series", [])
        hist = [p for p in series if p.get("kind") == "history"]
        fc = [p for p in series if p.get("kind") == "forecast"]
        if hist:
            _kv(doc, "Last observed", f"{hist[-1]['time'][:10]} = {hist[-1]['value']:.1f}")
        table_caption("Point forecast with approximate 95% prediction band.")
        ft = doc.add_table(rows=1, cols=3); ft.style = "Table Grid"
        for j, h in enumerate(["Period", "Forecast", "95% band"]):
            ft.rows[0].cells[j].text = h
        for p in fc:
            c = ft.add_row().cells
            c[0].text, c[1].text = p["time"][:10], f"{p['value']:.1f}"
            c[2].text = f"[{p['lower']:.1f}, {p['upper']:.1f}]"
        _footnote(doc, "The band widens with horizon and is residual-based (approximate, not exact).")
    elif mt in ("assurance", "interim"):
        v = m.get("verdict", {})
        if mt == "assurance":
            _kv(doc, "Assurance (probability of success)", f"{v.get('assurance', 0):.1%}")
            _kv(doc, "Power at the Target Value", f"{v.get('power', 0):.1%}")
        else:
            _kv(doc, "Predictive probability of success", f"{v.get('predictive_prob', 0):.1%}")
            if m.get("arms"):                              # two-arm: per-arm rates + the risk difference
                _kv(doc, "Posterior risk difference (t - c)", f"{v.get('posterior_diff', 0):+.1%} "
                    f"[{v.get('diff_ci_low', 0):+.1%}, {v.get('diff_ci_high', 0):+.1%}]")
                table_caption("Per-arm posterior response rates with 95% credible intervals.")
                pt = doc.add_table(rows=1, cols=4); pt.style = "Table Grid"
                for j, h in enumerate(["Arm", "Rate", "95% CrI", "n"]):
                    pt.rows[0].cells[j].text = h
                for a in m["arms"]:
                    c = pt.add_row().cells
                    tag = " (control)" if a.get("is_baseline") else ""
                    c[0].text = f"{a['arm']}{tag}"; c[1].text = f"{a['value']:.1%}"
                    c[2].text = f"[{a['ci_low']:.1%}, {a['ci_high']:.1%}]"; c[3].text = f"{a['n']:,}"
            else:
                _kv(doc, "Posterior response rate", f"{v.get('posterior_mean', 0):.1%}")
        rb = m.get("robustness") or {}
        if rb.get("panel"):
            table_caption("Prior sensitivity: the verdict under each defensible prior. A verdict that "
                          "flips across priors is prior-driven, not data-driven.")
            pt = doc.add_table(rows=1, cols=4); pt.style = "Table Grid"
            for j, h in enumerate(["Prior", "Parameters", "Assurance", "Verdict"]):
                pt.rows[0].cells[j].text = h
            for row in rb["panel"]:
                c = pt.add_row().cells
                c[0].text = str(row["prior"]); c[1].text = str(row["params"])
                c[2].text = f"{row['assurance']:.1%}"; c[3].text = str(row["call"])
            _footnote(doc, "FDA's Jan-2026 draft Bayesian guidance requires a prior-sensitivity "
                           "analysis. A FRAGILE verdict is reported as fragile, not as an answer.")
        if rb.get("type_i_error") is not None:
            table_caption("Operating characteristics implied by the pre-specified decision rule.")
            ot = doc.add_table(rows=1, cols=2); ot.style = "Table Grid"
            ot.rows[0].cells[0].text = "Quantity"; ot.rows[0].cells[1].text = "Value"
            for k, val in [("Type I error (GO rate at the LRV)", rb["type_i_error"]),
                           ("Power (GO rate at the TV)", rb["power"])]:
                c = ot.add_row().cells; c[0].text = k; c[1].text = f"{val:.1%}"
    elif m.get("arms"):
        binm = all(0 <= a["value"] <= 1 for a in m["arms"])
        table_caption("Per-arm summary with 95% confidence intervals.")
        at = doc.add_table(rows=1, cols=4); at.style = "Table Grid"
        for j, h in enumerate(["Arm", "Estimate", "95% CI", "n"]):
            at.rows[0].cells[j].text = h
        for a in m["arms"]:
            c = at.add_row().cells
            val = f"{a['value'] * 100:.1f}%" if binm else f"{a['value']:.3f}"
            ci = ("—" if a["ci_low"] != a["ci_low"] else
                  (f"[{a['ci_low'] * 100:.1f}%, {a['ci_high'] * 100:.1f}%]" if binm
                   else f"[{a['ci_low']:.3f}, {a['ci_high']:.3f}]"))
            c[0].text, c[1].text, c[2].text, c[3].text = str(a["arm"]), val, ci, f"{a['n']:,}"
        _footnote(doc, "Binary arms use Wilson score intervals; the between-arm difference uses the "
                       "Newcombe (proportions) or Welch (means) interval.")
    elif m.get("terms"):
        table_caption(f"Model estimates ({m.get('effect_label', 'estimate')}) with 95% CIs.")
        terms = m["terms"]
        has_n = any(t.get("n") is not None for t in terms)          # per-category subjects (categoricals)
        has_ev = any(t.get("events") is not None for t in terms)    # per-category events (event models)
        headers = (["Term"] + (["N"] if has_n else []) + (["Events"] if has_ev else [])
                   + [m.get("effect_label", "estimate"), "95% CI", "p-value"])
        tt = doc.add_table(rows=1, cols=len(headers)); tt.style = "Table Grid"
        for j, h in enumerate(headers):
            tt.rows[0].cells[j].text = h
        for term in terms:
            lo, pv = term["ci_low"], term["p"]
            vals = [str(term["name"])]
            if has_n:
                vals.append("—" if term.get("n") is None else f"{term['n']:,}")
            if has_ev:
                vals.append("—" if term.get("events") is None else f"{term['events']:,}")
            vals += [f"{term['estimate']:.3f}",
                     "—" if lo != lo else f"[{lo:.3f}, {term['ci_high']:.3f}]",
                     "—" if pv != pv else f"{pv:.4f}"]
            c = tt.add_row().cells
            for j, v in enumerate(vals):
                c[j].text = v
        _footnote(doc, "N and Events are the mutually-exclusive subjects and events per category (the "
                       "reference level is included). CIs are two-sided at 95%; p-values are unadjusted "
                       "unless the method note states otherwise.")
    elif getattr(result, "dataframe", None) is not None:
        full = result.dataframe
        df = full.iloc[:30, :12]                      # cap rows AND columns for a readable page
        table_caption("Query result (extract).")
        dt = doc.add_table(rows=1, cols=len(df.columns)); dt.style = "Table Grid"
        for j, col in enumerate(df.columns):
            dt.rows[0].cells[j].text = str(col)
        for _, row in df.iterrows():
            c = dt.add_row().cells
            for j, col in enumerate(df.columns):
                c[j].text = _cell(row[col])
        if full.shape[0] > 30 or full.shape[1] > 12:
            _footnote(doc, f"Showing {min(30, full.shape[0])} of {full.shape[0]:,} rows and "
                           f"{min(12, full.shape[1])} of {full.shape[1]} columns.")

    from agent import charts as _ch
    with _ch.render_for_print():                      # light/print palette so figures are legible on white
        for fig, caption in _figure_titles(result, m):
            png = _chart_png(fig)
            if png:
                figure_caption(caption)
                doc.add_picture(io.BytesIO(png), width=Inches(6.0))

    # ───────────────────────── 6. assumptions & diagnostics ─────────────────────────
    if diag:
        doc.add_heading("6. Assumptions and diagnostics", 1)
        for s in diag:
            doc.add_paragraph(s, style="List Bullet")
    if getattr(result, "findings", None):
        doc.add_heading("6.1 Statistical guardrail findings", 2)
        for f in result.findings:
            doc.add_paragraph(f"[{f.severity.upper()}] {f.kind} — {f.message}", style="List Bullet")

    # ───────────────────────── 7. interpretation ─────────────────────────
    if getattr(result, "interpretation", None):
        doc.add_heading("7. Interpretation and conclusions", 1)
        for line in result.interpretation.split("\n"):
            if line.strip():
                doc.add_paragraph(line.replace("**", "").lstrip("# ").strip())

    # ───────────────────────── 8. limitations & validation ─────────────────────────
    doc.add_heading("8. Limitations and validation statement", 1)
    doc.add_paragraph(
        "Data are synthetic and illustrative; magnitudes are not empirical. Variable selection here is "
        "data-driven (collinearity / quasi-constant / high-cardinality removal), which is appropriate "
        "for exploration but NOT for a confirmatory analysis, where the analysis set and methods must "
        "be pre-specified in a Statistical Analysis Plan (ICH E9). Confidence intervals and p-values are "
        "not adjusted for the model-selection process. This report is machine-generated and has not been "
        "independently double-programmed or source-data verified. A qualified biostatistician must "
        "review it before any decision or regulatory use.")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
